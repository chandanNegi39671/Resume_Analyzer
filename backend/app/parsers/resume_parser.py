from __future__ import annotations

import io
import re
from collections import Counter
from typing import Any

import pdfplumber
from docx import Document

PDF_MIME = "application/pdf"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

SECTION_PATTERNS = {
    "summary": re.compile(r"^\s*(professional\s+summary|summary|profile|objective)\s*$", re.I),
    "education": re.compile(r"^\s*(education|academic\s+background|qualifications?)\s*$", re.I),
    "experience": re.compile(
        r"^\s*(work\s+experience|professional\s+experience|experience|employment\s+history)\s*$",
        re.I,
    ),
    "skills": re.compile(r"^\s*(skills|technical\s+skills|core\s+competencies)\s*$", re.I),
    "projects": re.compile(r"^\s*(projects?|academic\s+projects?|personal\s+projects?)\s*$", re.I),
    "certifications": re.compile(r"^\s*(certifications?|licenses?)\s*$", re.I),
}


def _normalize_whitespace(value: str) -> str:
    value = value.replace("\u00a0", " ")
    value = re.sub(r"[ \t]+", " ", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def _extract_text_and_flags_from_pdf(file_bytes: bytes) -> tuple[str, list[str], list[str], list[str]]:
    page_texts: list[str] = []
    ats_flags: list[str] = []
    top_lines: list[str] = []
    bottom_lines: list[str] = []

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        if not pdf.pages:
            raise ValueError("Uploaded PDF has no pages.")

        table_detected = False
        images_detected = False
        multi_column_detected = False

        for page in pdf.pages:
            text = page.extract_text() or ""
            text = _normalize_whitespace(text)
            if text:
                page_texts.append(text)
                lines = [line.strip() for line in text.splitlines() if line.strip()]
                if lines:
                    top_lines.append(lines[0].lower())
                    bottom_lines.append(lines[-1].lower())

            tables = page.extract_tables()
            if tables:
                table_detected = True

            if getattr(page, "images", None):
                if len(page.images) > 0:
                    images_detected = True

            words = page.extract_words() or []
            if words:
                x_positions = sorted(float(word.get("x0", 0)) for word in words)
                if x_positions:
                    min_x = min(x_positions)
                    max_x = max(x_positions)
                    width = max_x - min_x
                    if width > 0:
                        left_count = sum(1 for x in x_positions if x <= min_x + width * 0.45)
                        right_count = sum(1 for x in x_positions if x >= min_x + width * 0.55)
                        center_count = sum(1 for x in x_positions if min_x + width * 0.45 < x < min_x + width * 0.55)
                        if left_count > 20 and right_count > 20 and center_count < 10:
                            multi_column_detected = True

        if table_detected:
            ats_flags.append("tables_present")
        if images_detected:
            ats_flags.append("images_present")
        if multi_column_detected:
            ats_flags.append("multi_column_layout")

    return "\n\n".join(page_texts), ats_flags, top_lines, bottom_lines


def _extract_text_and_flags_from_docx(file_bytes: bytes) -> tuple[str, list[str], list[str], list[str]]:
    ats_flags: list[str] = []
    doc = Document(io.BytesIO(file_bytes))

    body_lines = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text and paragraph.text.strip()]
    text = _normalize_whitespace("\n".join(body_lines))

    if doc.tables:
        ats_flags.append("tables_present")

    if hasattr(doc, "inline_shapes") and len(doc.inline_shapes) > 0:
        ats_flags.append("images_present")

    header_lines: list[str] = []
    footer_lines: list[str] = []
    for section in doc.sections:
        for p in section.header.paragraphs:
            if p.text and p.text.strip():
                header_lines.append(p.text.strip().lower())
        for p in section.footer.paragraphs:
            if p.text and p.text.strip():
                footer_lines.append(p.text.strip().lower())

    if header_lines or footer_lines:
        ats_flags.append("headers_footers_present")

    if "\t" in text:
        tab_count = text.count("\t")
        if tab_count >= 10:
            ats_flags.append("multi_column_layout")

    return text, ats_flags, header_lines, footer_lines


def _detect_non_standard_and_symbol_flags(raw_text: str, ats_flags: list[str]) -> None:
    if re.search(r"[^\x09\x0A\x0D\x20-\x7E]", raw_text):
        ats_flags.append("non_standard_characters")

    special_chars = re.findall(r"[^a-zA-Z0-9\s]", raw_text)
    if raw_text:
        ratio = len(special_chars) / max(len(raw_text), 1)
        if ratio > 0.08:
            ats_flags.append("excessive_special_symbols")


def _detect_header_footer_repetition(top_lines: list[str], bottom_lines: list[str], ats_flags: list[str]) -> None:
    repeated = False
    if len(top_lines) >= 2:
        common_top = Counter(top_lines).most_common(1)
        if common_top and common_top[0][1] >= 2:
            repeated = True
    if len(bottom_lines) >= 2:
        common_bottom = Counter(bottom_lines).most_common(1)
        if common_bottom and common_bottom[0][1] >= 2:
            repeated = True
    if repeated:
        ats_flags.append("headers_footers_present")


def _detect_sections(raw_text: str) -> dict[str, str]:
    lines = [line.strip() for line in raw_text.splitlines()]
    sections: dict[str, list[str]] = {
        "summary": [],
        "education": [],
        "experience": [],
        "skills": [],
        "projects": [],
        "certifications": [],
    }

    active_section = "summary"
    for line in lines:
        if not line:
            continue

        matched_section = None
        for section_name, pattern in SECTION_PATTERNS.items():
            if pattern.match(line):
                matched_section = section_name
                break

        if matched_section:
            active_section = matched_section
            continue

        sections[active_section].append(line)

    section_map = {key: "\n".join(value).strip() for key, value in sections.items()}
    return {
        "Summary": section_map["summary"],
        "Education": section_map["education"],
        "Experience": section_map["experience"],
        "Skills": section_map["skills"],
        "Projects": section_map["projects"],
        "Certifications": section_map["certifications"],
    }


def parse_resume(file_bytes: bytes, mime_type: str) -> dict[str, Any]:
    mime_type = (mime_type or "").strip().lower()
    if mime_type not in {PDF_MIME, DOCX_MIME}:
        raise ValueError("Unsupported file MIME type. Only PDF and DOCX are allowed.")

    if mime_type == PDF_MIME:
        raw_text, ats_flags, top_lines, bottom_lines = _extract_text_and_flags_from_pdf(file_bytes)
    else:
        raw_text, ats_flags, top_lines, bottom_lines = _extract_text_and_flags_from_docx(file_bytes)

    raw_text = _normalize_whitespace(raw_text)
    if not raw_text:
        raise ValueError("No readable text found in the uploaded file.")

    _detect_header_footer_repetition(top_lines, bottom_lines, ats_flags)
    _detect_non_standard_and_symbol_flags(raw_text, ats_flags)

    sections = _detect_sections(raw_text)
    word_count = len(re.findall(r"\b\w+\b", raw_text))

    return {
        "raw_text": raw_text,
        "sections": sections,
        "ats_flags": sorted(set(ats_flags)),
        "word_count": word_count,
    }
